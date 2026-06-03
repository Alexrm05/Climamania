from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
SOURCE_MD = BASE_DIR / "MANUAL_USUARIO_APP.md"
OUTPUT_DIR = BASE_DIR / "output" / "doc"
OUTPUT_DOCX = OUTPUT_DIR / "Manual_Usuario_ClimaMania.docx"


PLACEHOLDERS = {
    "2. Estructura general de la app": "insertar pantalla de estructura general de la app aqui",
    "11. Gestión de fotos, documentos y vídeos": "insertar pantalla del listado de fotos y documentos aqui",
    "14. Flujo completo del conforme del cliente": "insertar pantalla del flujo documental del conforme cliente aqui",
    "15. Detalle del PDF de conformidad": "insertar pantalla del pdf de conformidad aqui",
    "16. Detalle del PDF BOE": "insertar pantalla del pdf boe aqui",
    "31. Valoraciones": "insertar pantalla del qr de valoraciones aqui",
}


def set_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    fld_char3 = OxmlElement("w:t")
    fld_char3.text = "1"
    fld_char4 = OxmlElement("w:fldChar")
    fld_char4.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    run._r.append(fld_char3)
    run._r.append(fld_char4)


def parse_markdown(source: Path):
    lines = source.read_text(encoding="utf-8").splitlines()
    events = []
    paragraph = []

    def flush():
        if paragraph:
            text = " ".join(x.strip() for x in paragraph if x.strip()).strip()
            if text:
                events.append(("p", text))
            paragraph.clear()

    for line in lines:
        raw = line.rstrip()
        stripped = raw.strip()
        if not stripped:
            flush()
            continue
        if stripped == "---":
            flush()
            events.append(("hr",))
            continue
        if raw.startswith("# "):
            flush()
            events.append(("h1", raw[2:].strip()))
            continue
        if raw.startswith("#### "):
            flush()
            events.append(("h4", raw[5:].strip()))
            continue
        if raw.startswith("### "):
            flush()
            events.append(("h3", raw[4:].strip()))
            continue
        if raw.startswith("## "):
            flush()
            events.append(("h2", raw[3:].strip()))
            continue
        if re.match(r"^\s*[-*]\s+", raw):
            flush()
            indent = len(raw) - len(raw.lstrip(" "))
            text = re.sub(r"^\s*[-*]\s+", "", raw)
            events.append(("bullet", indent, text))
            continue
        if re.match(r"^\s*\d+\.\s+", raw):
            flush()
            indent = len(raw) - len(raw.lstrip(" "))
            text = raw.strip()
            events.append(("number", indent, text))
            continue
        paragraph.append(stripped)

    flush()
    return events


def add_inline_runs(paragraph, text: str):
    text = text.replace("\t", "    ")
    parts = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.bold = True
            run.font.name = "Courier New"
        else:
            paragraph.add_run(part)


def add_placeholder(document: Document, text: str):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"*{text}*")
    run.italic = True
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(34, 139, 34)


def build_docx():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(22)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(36, 49, 66)

    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(16)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(36, 49, 66)

    styles["Heading 3"].font.name = "Arial"
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.bold = True
    styles["Heading 3"].font.color.rgb = RGBColor(63, 78, 93)

    styles["Heading 4"].font.name = "Arial"
    styles["Heading 4"].font.size = Pt(10.5)
    styles["Heading 4"].font.bold = True
    styles["Heading 4"].font.italic = True
    styles["Heading 4"].font.color.rgb = RGBColor(89, 100, 116)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Manual completo de uso de la app ClimaMania")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(24)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(
        "Guia integral de navegacion, trabajo diario, documentacion, visitas, incidencias, presupuestos y cierre de instalaciones."
    )
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(107, 107, 107)

    version = document.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = version.add_run(f"Version generada el {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(120, 120, 120)

    document.add_paragraph("")
    add_placeholder(document, "insertar portada o pantalla principal de la app aqui")
    document.add_page_break()

    events = parse_markdown(SOURCE_MD)
    numbered_counters = {}

    for event in events:
        kind = event[0]
        if kind == "h1":
            p = document.add_paragraph(style="Heading 1")
            add_inline_runs(p, event[1])
        elif kind == "h2":
            heading = event[1]
            p = document.add_paragraph(style="Heading 2")
            add_inline_runs(p, heading)
            if heading in PLACEHOLDERS:
                add_placeholder(document, PLACEHOLDERS[heading])
        elif kind == "h3":
            p = document.add_paragraph(style="Heading 3")
            add_inline_runs(p, event[1])
        elif kind == "h4":
            p = document.add_paragraph(style="Heading 4")
            add_inline_runs(p, event[1])
        elif kind == "p":
            p = document.add_paragraph(style="Normal")
            p.paragraph_format.space_after = Pt(4)
            add_inline_runs(p, event[1])
        elif kind == "bullet":
            indent = int(event[1])
            p = document.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Cm(0.5 + (indent * 0.2))
            p.paragraph_format.space_after = Pt(2)
            p.style = document.styles["List Bullet"]
            add_inline_runs(p, event[2])
        elif kind == "number":
            indent = int(event[1])
            p = document.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Cm(0.5 + (indent * 0.2))
            p.paragraph_format.space_after = Pt(2)
            p.style = document.styles["List Number"]
            add_inline_runs(p, re.sub(r"^\d+\.\s*", "", event[2]))
        elif kind == "hr":
            document.add_paragraph("")

    # Footer with page number
    for sec in document.sections:
        footer = sec.footer
        para = footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run("Pagina ")
        run.font.name = "Arial"
        run.font.size = Pt(8)
        set_page_number(para)

    document.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_docx()
